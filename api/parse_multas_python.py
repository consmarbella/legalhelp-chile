import os
import io
import json
import requests
import re
from datetime import datetime
from dateutil.relativedelta import relativedelta
from flask import Flask, request, jsonify, send_file
import pdfplumber
from docx import Document

app = Flask(__name__)

DEEPSEEK_API_KEY = os.environ.get("DEEPSEEK_API_KEY", "")

SYSTEM_PROMPT = """Actúa como un extractor de datos legales chileno ultra-preciso. Tu único objetivo es extraer las multas de tránsito/TAG del texto provisto y estructurarlas en un arreglo JSON. Ignora cualquier texto introductorio. Si una multa no tiene Juzgado o Fecha de Anotación (también llamada FECHA INGRESO RMNP), ignórala. Asegúrate de extraer bien la patente, el juzgado, la fecha y el rol.

Formato de Salida Esperado:
{
  "multas": [
    {
      "juzgado": "1ER JUZGADO DE POLICIA LOCAL DE PUDAHUEL",
      "fecha_anotacion": "2022-05-14",
      "causa_rol": "12345-2022",
      "patente": "ABCD12",
      "monto_utm": 1.0
    }
  ]
}"""

def sanitize_jpl(raw_jpl):
    jpl = raw_jpl.upper().strip()
    # Replace common variations
    jpl = re.sub(r'\b1ER\b', '1°', jpl)
    jpl = re.sub(r'\b1\s+JUZGADO\b', '1° JUZGADO', jpl)
    jpl = re.sub(r'\b2DO\b', '2°', jpl)
    jpl = re.sub(r'\b2\s+JUZGADO\b', '2° JUZGADO', jpl)
    jpl = re.sub(r'\b3ER\b', '3°', jpl)
    jpl = re.sub(r'\b3\s+JUZGADO\b', '3° JUZGADO', jpl)
    # Remove accents for grouping
    jpl = jpl.replace('Á', 'A').replace('É', 'E').replace('Í', 'I').replace('Ó', 'O').replace('Ú', 'U')
    return jpl

@app.route('/api/parse_multas_python', methods=['POST'])
def parse_multas():
    if 'file' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
        
    file = request.files['file']
    
    # Datos desde el formdata
    patente_form = request.form.get('patente', '')
    rutSolicitante = request.form.get('rutSolicitante', '')
    nombreSolicitante = request.form.get('nombreSolicitante', '')
    
    try:
        filename = file.filename.lower()
        file_bytes = file.read()
        full_text = ""
        
        if filename.endswith('.pdf'):
            # Load PDF using pdfplumber
            try:
                pdf = pdfplumber.open(io.BytesIO(file_bytes))
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        full_text += text + "\n"
            except Exception as e:
                return jsonify({"error": "Error leyendo PDF. Asegúrese de subir un documento digital válido."}), 400
        else:
            return jsonify({"error": "Solo se aceptan archivos PDF."}), 400

        if not full_text.strip():
            return jsonify({"error": "El PDF no contiene texto legible (parece una imagen escaneada)."}), 400

        # LLM Extraction
        if not DEEPSEEK_API_KEY:
            return jsonify({"error": "DEEPSEEK_API_KEY no está configurada en el backend."}), 500

        response = requests.post(
            'https://api.deepseek.com/chat/completions',
            headers={
                'Content-Type': 'application/json',
                'Authorization': f'Bearer {DEEPSEEK_API_KEY}'
            },
            json={
                "model": "deepseek-chat",
                "messages": [
                    {"role": "system", "content": SYSTEM_PROMPT},
                    {"role": "user", "content": full_text[:15000]} # Limit text to avoid token limits
                ],
                "response_format": {"type": "json_object"},
                "temperature": 0.1
            },
            timeout=30
        )
        
        if response.status_code != 200:
            return jsonify({"error": "Error al comunicarse con DeepSeek", "details": response.text}), 502
            
        data = response.json()
        content = data.get('choices', [{}])[0].get('message', {}).get('content', '{}')
        
        try:
            parsed_json = json.loads(content)
        except json.JSONDecodeError:
            return jsonify({"error": "DeepSeek no retornó un JSON válido."}), 502

        multas_extraidas = parsed_json.get('multas', [])
        
        # Física de fechas y agrupación
        fecha_actual = datetime.now()
        limite_prescripcion = fecha_actual - relativedelta(years=3)
        
        multas_elegibles = []
        global_patente = patente_form
        
        for multa in multas_extraidas:
            fecha_str = multa.get('fecha_anotacion', '')
            try:
                if '-' in fecha_str:
                    fecha_multa = datetime.strptime(fecha_str, "%Y-%m-%d")
                elif '/' in fecha_str:
                    fecha_multa = datetime.strptime(fecha_str, "%d/%m/%Y")
                else:
                    continue
            except ValueError:
                continue
                
            if fecha_multa < limite_prescripcion:
                if multa.get('patente') and not global_patente:
                    global_patente = multa['patente']
                multas_elegibles.append(multa)
                
        # Extract owner details fallback
        nombre_extracted = nombreSolicitante
        run_extracted = rutSolicitante
        
        nombre_match = re.search(r'Nombre\s*:\s*(.+)', full_text, re.IGNORECASE)
        if nombre_match and not nombre_extracted:
            nombre_extracted = nombre_match.group(1).strip()
            
        run_match = re.search(r'R\.U\.N\.\s*:\s*([0-9\.\-kK]+)', full_text, re.IGNORECASE)
        if run_match and not run_extracted:
            run_extracted = run_match.group(1).strip()
            
        # Group by comuna/juzgado
        multas_por_juzgado = {}
        comunas_con_correo = []
        
        for m in multas_elegibles:
            jpl = sanitize_jpl(m.get('juzgado', 'DESCONOCIDO'))
            if jpl not in multas_por_juzgado:
                multas_por_juzgado[jpl] = []
            
            # Formatear para el frontend
            multas_por_juzgado[jpl].append({
                "id": m.get('causa_rol', 'S/R'),
                "fecha_ingreso": m.get('fecha_anotacion', ''),
                "rol": m.get('causa_rol', 'S/R'),
                "patente": m.get('patente', global_patente),
                "monto_utm": m.get('monto_utm', 0)
            })

        for tribunal, lista in multas_por_juzgado.items():
            comunas_con_correo.append({
                "nombre": tribunal,
                "correo": f"jpl@{tribunal.lower().replace(' ', '').replace('°', '')}.cl",
                "multas": lista
            })
            
        # Cobro
        total_cobro = 0 if not multas_por_juzgado else 10000 + max(0, len(multas_por_juzgado) - 1) * 4000
        
        return jsonify({
            "ok": True,
            "cliente": {
                "nombre": nombre_extracted,
                "run": run_extracted,
                "patente": global_patente
            },
            "multas_por_tribunal": multas_por_juzgado,
            "totalMultas": len(multas_elegibles),
            "comunas": comunas_con_correo,
            "cobro": total_cobro
        })

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

@app.route('/api/generate_template', methods=['POST'])
def generate_template():
    try:
        data = request.get_json()
        if not data:
            return jsonify({'ok': False, 'error': 'No JSON payload provided'}), 400
            
        template_id = data.get('template_id')
        variables = data.get('variables', {})
        
        if not template_id:
            return jsonify({'ok': False, 'error': 'template_id is required'}), 400
            
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        template_path = os.path.join(base_dir, 'plantillas', f"{template_id}.txt")
        
        if not os.path.exists(template_path):
            template_path = os.path.join(os.getcwd(), 'plantillas', f"{template_id}.txt")
            if not os.path.exists(template_path):
                return jsonify({'ok': False, 'error': f'Plantilla {template_id} no encontrada'}), 404
                
        with open(template_path, 'r', encoding='utf-8') as f:
            template_text = f.read()
            
        for key, value in variables.items():
            if value is None:
                value = ""
            placeholder = "{{" + str(key) + "}}"
            template_text = template_text.replace(placeholder, str(value))
            
        return jsonify({
            'ok': True,
            'documento': template_text,
            'fuente': 'python-deterministic-template'
        })
        
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500

@app.route('/api/generate_tag_zip', methods=['POST'])
def generate_tag_zip():
    try:
        data = request.get_json()
        if not data:
            return jsonify({'ok': False, 'error': 'No JSON payload provided'}), 400
            
        variables = data.get('variables', {})
        multas_por_tribunal = data.get('multas_por_tribunal', {})
        
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        template_path = os.path.join(base_dir, 'plantillas', "jpl-prescripcion-tag.txt")
        
        if not os.path.exists(template_path):
            template_path = os.path.join(os.getcwd(), 'plantillas', "jpl-prescripcion-tag.txt")
            
        with open(template_path, 'r', encoding='utf-8') as f:
            base_template_text = f.read()

        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            for tribunal, multas in multas_por_tribunal.items():
                # Preparar texto específico para este tribunal
                texto = base_template_text
                
                # Inyectar variables del cliente
                cliente_vars = {
                    'TRIBUNAL': tribunal,
                    'NOMBRE_COMPLETO': variables.get('nombre', ''),
                    'RUT': variables.get('rut', ''),
                    'PROFESION': variables.get('actividad', 'Empleado'),
                    'DOMICILIO': variables.get('direccion', ''),
                    'COMUNA_DOMICILIO': variables.get('comunaDireccion', ''),
                    'PATENTE': variables.get('patente', ''),
                }
                
                for k, v in cliente_vars.items():
                    texto = texto.replace("{{" + str(k) + "}}", str(v))
                    
                # Crear documento Word
                doc = Document()
                # Insertar el texto generado
                for paragraph in texto.split('\n'):
                    doc.add_paragraph(paragraph)
                    
                # Añadir la tabla de multas determinista
                doc.add_paragraph("MULTAS ELEGIBLES PARA PRESCRIPCIÓN:")
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Rol'
                hdr_cells[1].text = 'Fecha Anotación'
                hdr_cells[2].text = 'Patente'
                
                for m in multas:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(m.get('rol', ''))
                    row_cells[1].text = str(m.get('fecha_ingreso', ''))
                    row_cells[2].text = str(m.get('patente', ''))
                    
                doc.add_paragraph("\nPOR TANTO,\nRUEGO A US. acceder a lo solicitado y declarar la prescripción.")
                
                doc_buffer = io.BytesIO()
                doc.save(doc_buffer)
                
                safe_tribunal_name = re.sub(r'[^a-zA-Z0-9_]', '_', tribunal)
                zip_file.writestr(f"Prescripcion_{safe_tribunal_name}.docx", doc_buffer.getvalue())

        zip_buffer.seek(0)
        return send_file(
            zip_buffer,
            mimetype='application/zip',
            as_attachment=True,
            download_name='Escritos_Legales_LegalHelp.zip'
        )

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'ok': False, 'error': str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True, port=5328)
